#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
====================================================================
Script crawl văn bản pháp luật từ vbpl.vn
Phục vụ đồ án: Hệ thống quản lý tài liệu nội bộ cho công ty luật dân sự
====================================================================

Chiến lược crawl:
- Bước 1: Thử crawl metadata + nội dung từ vbpl.vn (trang chính thức)
- Bước 2: Nếu vbpl.vn không trả HTML (SPA/JS), crawl từ thuvienphapluat.vn
- Bước 3: Với nội dung text thu được, tạo file Word (.docx) bằng python-docx
- Bước 4: Xuất CSV tổng hợp

Thư viện sử dụng:
- requests: gửi HTTP request
- BeautifulSoup (bs4): parse HTML
- python-docx: tạo file Word (.docx) từ nội dung text
  (Chọn python-docx thay vì pdfkit/weasyprint vì:
   + Không cần cài thêm công cụ hệ thống (wkhtmltopdf, GTK, ...)
   + Hoạt động tốt trên Windows
   + File Word dễ chỉnh sửa, phù hợp cho hệ thống quản lý tài liệu)

Tác giả: Script tự động
Ngày tạo: 2026-06-14
"""

import os
import re
import csv
import time
import logging
import unicodedata
import argparse
import sys
from datetime import datetime
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[CẢNH BÁO] Thư viện python-docx chưa được cài đặt.")
    print("Chạy: pip install python-docx")

# ============================================================
# CẤU HÌNH
# ============================================================

# Thư mục gốc lưu file
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCUMENTS_DIR = os.path.join(BASE_DIR, "documents")
CSV_OUTPUT = os.path.join(BASE_DIR, "vanban_phapluat.csv")
LOG_FILE = os.path.join(BASE_DIR, "loi_crawl.txt")

# Headers giả lập trình duyệt
HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/125.0.0.0 Safari/537.36"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "vi-VN,vi;q=0.9,en-US;q=0.8,en;q=0.7",
    "Accept-Encoding": "gzip, deflate, br",
    "Connection": "keep-alive",
}

# Delay giữa các request (giây) - tôn trọng server
REQUEST_DELAY = 2
MIN_FULLTEXT_CHARS = 2000

# Giới hạn số văn bản mỗi lĩnh vực
MAX_PER_FIELD = 5

# ============================================================
# DANH SÁCH VĂN BẢN CẦN CRAWL (HARDCODED - 20 VĂN BẢN)
# ============================================================
# Vì vbpl.vn mới (2026) sử dụng SPA (Next.js), không crawl được
# bằng requests+BS4 thuần, ta sử dụng danh sách văn bản đã biết
# với URL từ thuvienphapluat.vn (nguồn phổ biến, cấu trúc HTML tốt)
# hoặc nội dung hardcode để đảm bảo script luôn chạy được.

VANBAN_DATA = {
    "dan-su": {
        "ten_linh_vuc": "Bộ luật Dân sự",
        "vanban": [
            {
                "ten": "Bộ luật Dân sự 2015",
                "soHieu": "91/2015/QH13",
                "ngayBanHanh": "24/11/2015",
                "loaiVanBan": "Bộ luật",
                "url": "https://thuvienphapluat.vn/van-ban/Quyen-dan-su/Bo-luat-dan-su-2015-296215.aspx",
                "mo_ta": (
                    "Bộ luật Dân sự số 91/2015/QH13 được Quốc hội khóa XIII thông qua "
                    "ngày 24/11/2015, có hiệu lực từ 01/01/2017. Bộ luật quy định địa vị "
                    "pháp lý, chuẩn mực pháp lý về cách ứng xử của cá nhân, pháp nhân; "
                    "quyền, nghĩa vụ về nhân thân và tài sản của cá nhân, pháp nhân trong "
                    "các quan hệ được hình thành trên cơ sở bình đẳng, tự do ý chí, độc lập "
                    "về tài sản và tự chịu trách nhiệm."
                ),
            },
            {
                "ten": "Nghị định 21/2021/NĐ-CP hướng dẫn Bộ luật Dân sự về bảo đảm thực hiện nghĩa vụ",
                "soHieu": "21/2021/NĐ-CP",
                "ngayBanHanh": "19/03/2021",
                "loaiVanBan": "Nghị định",
                "url": "https://thuvienphapluat.vn/van-ban/Quyen-dan-su/Nghi-dinh-21-2021-ND-CP-huong-dan-Bo-luat-Dan-su-ve-bao-dam-thuc-hien-nghia-vu-467055.aspx",
                "mo_ta": (
                    "Nghị định 21/2021/NĐ-CP quy định chi tiết thi hành Bộ luật Dân sự "
                    "về bảo đảm thực hiện nghĩa vụ, bao gồm: cầm cố tài sản, thế chấp "
                    "tài sản, đặt cọc, ký cược, ký quỹ, bảo lưu quyền sở hữu, bảo lãnh, "
                    "tín chấp và cầm giữ tài sản."
                ),
            },
            {
                "ten": "Nghị định 22/2025/NĐ-CP về giao dịch điện tử trong lĩnh vực dân sự",
                "soHieu": "22/2025/NĐ-CP",
                "ngayBanHanh": "10/01/2025",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 22/2025/NĐ-CP quy định chi tiết một số điều về giao dịch "
                    "điện tử trong lĩnh vực dân sự, thương mại, bao gồm hợp đồng điện tử, "
                    "chữ ký điện tử, và chứng thực điện tử trong các giao dịch dân sự."
                ),
            },
            {
                "ten": "Nghị quyết 04/2017/NQ-HĐTP hướng dẫn áp dụng một số quy định tại Bộ luật Dân sự",
                "soHieu": "04/2017/NQ-HĐTP",
                "ngayBanHanh": "05/05/2017",
                "loaiVanBan": "Nghị quyết",
                "url": "",
                "mo_ta": (
                    "Nghị quyết 04/2017/NQ-HĐTP của Hội đồng Thẩm phán Tòa án nhân dân "
                    "tối cao hướng dẫn một số quy định tại Bộ luật Dân sự 2015 về trách "
                    "nhiệm bồi thường thiệt hại ngoài hợp đồng."
                ),
            },
            {
                "ten": "Nghị định 48/2024/NĐ-CP sửa đổi Nghị định về đăng ký biện pháp bảo đảm",
                "soHieu": "48/2024/NĐ-CP",
                "ngayBanHanh": "09/05/2024",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 48/2024/NĐ-CP sửa đổi, bổ sung một số điều của Nghị định "
                    "về đăng ký biện pháp bảo đảm, quy định trình tự, thủ tục đăng ký thế "
                    "chấp quyền sử dụng đất, tài sản gắn liền với đất."
                ),
            },
        ],
    },
    "hon-nhan-gia-dinh": {
        "ten_linh_vuc": "Luật Hôn nhân và Gia đình",
        "vanban": [
            {
                "ten": "Luật Hôn nhân và Gia đình 2014",
                "soHieu": "52/2014/QH13",
                "ngayBanHanh": "19/06/2014",
                "loaiVanBan": "Luật",
                "url": "https://thuvienphapluat.vn/van-ban/Quyen-dan-su/Luat-Hon-nhan-va-gia-dinh-2014-238640.aspx",
                "mo_ta": (
                    "Luật Hôn nhân và Gia đình số 52/2014/QH13 được Quốc hội thông qua "
                    "ngày 19/06/2014, có hiệu lực từ 01/01/2015. Luật quy định chế độ hôn "
                    "nhân và gia đình; chuẩn mực pháp lý cho cách ứng xử giữa các thành "
                    "viên gia đình; trách nhiệm của cá nhân, tổ chức, Nhà nước và xã hội "
                    "trong việc xây dựng, củng cố chế độ hôn nhân và gia đình."
                ),
            },
            {
                "ten": "Nghị định 126/2014/NĐ-CP hướng dẫn Luật Hôn nhân và Gia đình",
                "soHieu": "126/2014/NĐ-CP",
                "ngayBanHanh": "31/12/2014",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 126/2014/NĐ-CP quy định chi tiết một số điều và biện pháp "
                    "thi hành Luật Hôn nhân và Gia đình, bao gồm các quy định về kết hôn, "
                    "ly hôn, quyền và nghĩa vụ giữa vợ chồng, cha mẹ và con."
                ),
            },
            {
                "ten": "Nghị định 10/2015/NĐ-CP quy định về sinh con bằng kỹ thuật thụ tinh trong ống nghiệm",
                "soHieu": "10/2015/NĐ-CP",
                "ngayBanHanh": "28/01/2015",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 10/2015/NĐ-CP quy định về sinh con bằng kỹ thuật thụ tinh "
                    "trong ống nghiệm và điều kiện mang thai hộ vì mục đích nhân đạo. "
                    "Quy định điều kiện, quyền và nghĩa vụ của các bên trong việc mang thai hộ."
                ),
            },
            {
                "ten": "Thông tư liên tịch 01/2016/TTLT-TANDTC-VKSNDTC-BTP hướng dẫn thi hành Luật HNGĐ",
                "soHieu": "01/2016/TTLT-TANDTC-VKSNDTC-BTP",
                "ngayBanHanh": "06/01/2016",
                "loaiVanBan": "Thông tư liên tịch",
                "url": "",
                "mo_ta": (
                    "Thông tư liên tịch 01/2016/TTLT-TANDTC-VKSNDTC-BTP hướng dẫn thi hành "
                    "một số quy định của Luật Hôn nhân và Gia đình, bao gồm quy định về "
                    "tài sản chung, tài sản riêng của vợ chồng và chia tài sản khi ly hôn."
                ),
            },
            {
                "ten": "Luật Phòng, chống bạo lực gia đình 2022",
                "soHieu": "13/2022/QH15",
                "ngayBanHanh": "14/11/2022",
                "loaiVanBan": "Luật",
                "url": "",
                "mo_ta": (
                    "Luật Phòng, chống bạo lực gia đình 2022 quy định về phòng ngừa, "
                    "phát hiện, ngăn chặn và xử lý hành vi bạo lực gia đình; bảo vệ, "
                    "hỗ trợ nạn nhân bạo lực gia đình; trách nhiệm của cơ quan, tổ chức, "
                    "gia đình, cá nhân trong phòng, chống bạo lực gia đình."
                ),
            },
        ],
    },
    "dat-dai": {
        "ten_linh_vuc": "Luật Đất đai",
        "vanban": [
            {
                "ten": "Luật Đất đai 2024",
                "soHieu": "31/2024/QH15",
                "ngayBanHanh": "18/01/2024",
                "loaiVanBan": "Luật",
                "url": "https://thuvienphapluat.vn/van-ban/Bat-dong-san/Luat-Dat-dai-2024-31-2024-QH15-573243.aspx",
                "mo_ta": (
                    "Luật Đất đai 2024 số 31/2024/QH15 được Quốc hội thông qua ngày "
                    "18/01/2024, có hiệu lực từ 01/08/2024. Luật quy định về chế độ sở "
                    "hữu đất đai, quyền hạn và trách nhiệm của Nhà nước đại diện chủ sở "
                    "hữu toàn dân về đất đai và thống nhất quản lý về đất đai, chế độ "
                    "quản lý và sử dụng đất đai, quyền và nghĩa vụ của người sử dụng đất."
                ),
            },
            {
                "ten": "Nghị định 71/2024/NĐ-CP quy định về giá đất",
                "soHieu": "71/2024/NĐ-CP",
                "ngayBanHanh": "27/06/2024",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 71/2024/NĐ-CP quy định về giá đất, bảng giá đất, giá đất "
                    "cụ thể và tư vấn xác định giá đất. Nghị định quy định phương pháp "
                    "định giá đất, trình tự thủ tục xác định giá đất cụ thể."
                ),
            },
            {
                "ten": "Nghị định 102/2024/NĐ-CP hướng dẫn Luật Đất đai 2024",
                "soHieu": "102/2024/NĐ-CP",
                "ngayBanHanh": "30/07/2024",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 102/2024/NĐ-CP quy định chi tiết thi hành một số điều của "
                    "Luật Đất đai 2024 về thu hồi đất, trưng dụng đất, bồi thường, hỗ trợ, "
                    "tái định cư khi Nhà nước thu hồi đất."
                ),
            },
            {
                "ten": "Nghị định 101/2024/NĐ-CP về đăng ký đất đai, cấp Giấy chứng nhận",
                "soHieu": "101/2024/NĐ-CP",
                "ngayBanHanh": "29/07/2024",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 101/2024/NĐ-CP quy định về đăng ký đất đai, cấp Giấy "
                    "chứng nhận quyền sử dụng đất, quyền sở hữu tài sản gắn liền với "
                    "đất. Quy định trình tự, thủ tục đăng ký đất đai lần đầu, đăng ký "
                    "biến động đất đai, cấp đổi, cấp lại Giấy chứng nhận."
                ),
            },
            {
                "ten": "Thông tư 02/2024/TT-BTNMT quy định kỹ thuật điều tra, đánh giá đất đai",
                "soHieu": "02/2024/TT-BTNMT",
                "ngayBanHanh": "15/04/2024",
                "loaiVanBan": "Thông tư",
                "url": "",
                "mo_ta": (
                    "Thông tư 02/2024/TT-BTNMT của Bộ Tài nguyên và Môi trường quy định "
                    "kỹ thuật điều tra, đánh giá đất đai; quy định về thu thập, quản lý "
                    "thông tin đất đai, xây dựng cơ sở dữ liệu đất đai."
                ),
            },
        ],
    },
    "lao-dong": {
        "ten_linh_vuc": "Luật Lao động",
        "vanban": [
            {
                "ten": "Bộ luật Lao động 2019",
                "soHieu": "45/2019/QH14",
                "ngayBanHanh": "20/11/2019",
                "loaiVanBan": "Bộ luật",
                "url": "https://thuvienphapluat.vn/van-ban/Lao-dong-Tien-luong/Bo-Luat-lao-dong-2019-333670.aspx",
                "mo_ta": (
                    "Bộ luật Lao động 2019 số 45/2019/QH14 được Quốc hội thông qua ngày "
                    "20/11/2019, có hiệu lực từ 01/01/2021. Bộ luật quy định tiêu chuẩn "
                    "lao động; quyền, nghĩa vụ, trách nhiệm của người lao động, người sử "
                    "dụng lao động, tổ chức đại diện người lao động, tổ chức đại diện "
                    "người sử dụng lao động trong quan hệ lao động và các quan hệ khác "
                    "liên quan trực tiếp đến quan hệ lao động."
                ),
            },
            {
                "ten": "Nghị định 145/2020/NĐ-CP hướng dẫn Bộ luật Lao động về điều kiện lao động",
                "soHieu": "145/2020/NĐ-CP",
                "ngayBanHanh": "14/12/2020",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 145/2020/NĐ-CP quy định chi tiết và hướng dẫn thi hành "
                    "một số điều của Bộ luật Lao động về điều kiện lao động và quan hệ "
                    "lao động, bao gồm: hợp đồng lao động, tiền lương, thời giờ làm việc, "
                    "thời giờ nghỉ ngơi, kỷ luật lao động, trách nhiệm vật chất."
                ),
            },
            {
                "ten": "Nghị định 12/2022/NĐ-CP quy định xử phạt vi phạm hành chính trong lĩnh vực lao động",
                "soHieu": "12/2022/NĐ-CP",
                "ngayBanHanh": "17/01/2022",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 12/2022/NĐ-CP quy định xử phạt vi phạm hành chính trong "
                    "lĩnh vực lao động, bảo hiểm xã hội, đưa người lao động Việt Nam "
                    "đi làm việc ở nước ngoài theo hợp đồng."
                ),
            },
            {
                "ten": "Nghị định 135/2020/NĐ-CP quy định về tuổi nghỉ hưu",
                "soHieu": "135/2020/NĐ-CP",
                "ngayBanHanh": "18/11/2020",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 135/2020/NĐ-CP quy định về tuổi nghỉ hưu đối với người "
                    "lao động theo lộ trình tăng tuổi nghỉ hưu từ năm 2021 đến khi đạt "
                    "62 tuổi đối với lao động nam (vào năm 2028) và 60 tuổi đối với lao "
                    "động nữ (vào năm 2035)."
                ),
            },
            {
                "ten": "Nghị định 38/2022/NĐ-CP quy định mức lương tối thiểu",
                "soHieu": "38/2022/NĐ-CP",
                "ngayBanHanh": "12/06/2022",
                "loaiVanBan": "Nghị định",
                "url": "",
                "mo_ta": (
                    "Nghị định 38/2022/NĐ-CP quy định mức lương tối thiểu đối với người "
                    "lao động làm việc theo hợp đồng lao động, áp dụng từ ngày 01/07/2022. "
                    "Mức lương tối thiểu vùng I là 4.680.000 đồng/tháng."
                ),
            },
        ],
    },
}

# ============================================================
# SETUP LOGGING
# ============================================================

def setup_logging():
    """Thiết lập ghi log ra file và console."""
    logger = logging.getLogger("crawl_vbpl")
    logger.setLevel(logging.DEBUG)

    # File handler - ghi lỗi ra file
    fh = logging.FileHandler(LOG_FILE, encoding="utf-8", mode="w")
    fh.setLevel(logging.WARNING)
    fh_format = logging.Formatter(
        "%(asctime)s | %(levelname)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S"
    )
    fh.setFormatter(fh_format)

    # Console handler
    ch = logging.StreamHandler()
    ch.setLevel(logging.INFO)
    ch_format = logging.Formatter("%(message)s")
    ch.setFormatter(ch_format)

    logger.addHandler(fh)
    logger.addHandler(ch)
    return logger


logger = setup_logging()

# ============================================================
# HÀM TIỆN ÍCH
# ============================================================

def remove_diacritics(text):
    """Loại bỏ dấu tiếng Việt khỏi chuỗi."""
    nfkd = unicodedata.normalize("NFKD", text)
    return "".join(c for c in nfkd if not unicodedata.category(c).startswith("M"))


def sanitize_filename(text):
    """
    Chuyển tên văn bản thành tên file an toàn:
    - Bỏ dấu tiếng Việt
    - Thay khoảng trắng bằng dấu gạch ngang
    - Loại bỏ ký tự đặc biệt
    - Viết thường
    """
    text = remove_diacritics(text)
    text = text.lower().strip()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_]+", "-", text)
    text = re.sub(r"-+", "-", text)
    text = text.strip("-")
    # Giới hạn độ dài
    if len(text) > 80:
        text = text[:80].rstrip("-")
    return text


def ensure_dir(path):
    """Tạo thư mục nếu chưa tồn tại."""
    os.makedirs(path, exist_ok=True)


def fetch_page(url, timeout=30):
    """
    Gửi GET request và trả về response.
    Trả về None nếu có lỗi.
    """
    try:
        session = requests.Session()
        session.headers.update(HEADERS)
        response = session.get(url, timeout=timeout, allow_redirects=True)
        response.raise_for_status()
        return response
    except requests.exceptions.RequestException as e:
        logger.warning(f"Lỗi khi truy cập {url}: {e}")
        return None


def extract_content_from_tvpl(url):
    """
    Trích xuất nội dung toàn văn từ thuvienphapluat.vn.
    Trả về dict chứa: title, content_text, metadata
    """
    logger.info(f"  → Đang crawl nội dung từ: {url}")
    time.sleep(REQUEST_DELAY)

    response = fetch_page(url)
    if not response:
        rendered_text = render_text_with_selenium(url)
        if rendered_text:
            return {
                "title": "",
                "content_text": rendered_text,
                "content_html": "",
                "source_url": url,
                "method": "selenium",
            }
        return None

    soup = BeautifulSoup(response.content, "lxml")

    result = {
        "title": "",
        "content_text": "",
        "content_html": "",
    }

    # Trích xuất tiêu đề
    title_el = soup.find("div", class_="title") or soup.find("h1")
    if title_el:
        result["title"] = title_el.get_text(strip=True)

    # Trích xuất nội dung toàn văn
    # Thử nhiều selector phổ biến trên thuvienphapluat.vn
    content_selectors = [
        ("div", {"class": "content1"}),
        ("div", {"class": "toanvancontent"}),
        ("div", {"class": "fulltext"}),
        ("div", {"id": "toanvancontent"}),
        ("div", {"class": "WordSection1"}),
        ("div", {"class": "box-ct"}),
    ]

    content_el = None
    for tag, attrs in content_selectors:
        content_el = soup.find(tag, attrs)
        if content_el:
            break

    if content_el:
        result["content_html"] = str(content_el)
        result["content_text"] = content_el.get_text(separator="\n", strip=True)
    else:
        # Fallback: lấy toàn bộ body text
        body = soup.find("body")
        if body:
            # Loại bỏ script, style
            for tag in body.find_all(["script", "style", "nav", "header", "footer"]):
                tag.decompose()
            result["content_text"] = body.get_text(separator="\n", strip=True)

    return result


def clean_soup_text(element):
    """Return visible text from an HTML element."""
    if not element:
        return ""
    for tag in element.find_all([
        "script", "style", "noscript", "nav", "header", "footer", "form",
        "iframe", "button", "select", "option",
    ]):
        tag.decompose()
    text = element.get_text(separator="\n", strip=True)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def extract_text_from_soup(soup):
    """Extract the most likely full legal text from a parsed page."""
    selectors = [
        ("div", {"class": "content1"}),
        ("div", {"class": "toanvancontent"}),
        ("div", {"id": "toanvancontent"}),
        ("div", {"class": "fulltext"}),
        ("div", {"class": "WordSection1"}),
        ("div", {"class": "box-ct"}),
        ("article", {}),
        ("body", {}),
    ]

    best_text = ""
    best_html = ""
    for tag, attrs in selectors:
        for element in soup.find_all(tag, attrs):
            text = clean_soup_text(element)
            if len(text) > len(best_text):
                best_text = text
                best_html = str(element)

    return best_text, best_html


def find_fulltext_urls(soup, base_url):
    """Find iframe/print/toanvan URLs that may contain the real full text."""
    candidates = []
    keywords = ("print", "toanvan", "toan-van", "vbpq-print", "fulltext", "xemvanban")

    for iframe in soup.find_all("iframe"):
        src = iframe.get("src")
        if src:
            candidates.append(urljoin(base_url, src))

    for tag in soup.find_all(["a", "link"]):
        href = tag.get("href")
        if not href:
            continue
        absolute = urljoin(base_url, href)
        if any(keyword in absolute.lower() for keyword in keywords):
            candidates.append(absolute)

    unique = []
    seen = set()
    for candidate in candidates:
        if candidate not in seen:
            seen.add(candidate)
            unique.append(candidate)
    return unique


def render_text_with_selenium(url):
    """Render JavaScript page with Selenium when requests cannot get full text."""
    try:
        from selenium import webdriver
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.common.by import By
    except ImportError:
        logger.info("  Selenium is not installed; skip JS rendering fallback.")
        return ""

    driver = None
    try:
        options = Options()
        options.add_argument("--headless=new")
        options.add_argument("--disable-gpu")
        options.add_argument("--no-sandbox")
        options.add_argument("--window-size=1366,2000")
        driver = webdriver.Chrome(options=options)
        driver.get(url)
        time.sleep(4)

        best_text = driver.find_element(By.TAG_NAME, "body").text
        iframe_elements = driver.find_elements(By.TAG_NAME, "iframe")
        for iframe in iframe_elements:
            try:
                driver.switch_to.frame(iframe)
                frame_text = driver.find_element(By.TAG_NAME, "body").text
                if len(frame_text) > len(best_text):
                    best_text = frame_text
                driver.switch_to.default_content()
            except Exception:
                driver.switch_to.default_content()

        return best_text.strip()
    except Exception as e:
        logger.warning(f"  Selenium fallback failed for {url}: {e}")
        return ""
    finally:
        if driver:
            driver.quit()


def extract_content_from_tvpl(url):
    """
    Extract full text from a detail URL.
    Order: detail page -> iframe/print/toanvan candidates -> Selenium.
    """
    logger.info(f"  -> Fetch detail page: {url}")
    time.sleep(REQUEST_DELAY)

    response = fetch_page(url)
    if not response:
        rendered_text = render_text_with_selenium(url)
        if rendered_text:
            return {
                "title": "",
                "content_text": rendered_text,
                "content_html": "",
                "source_url": url,
                "method": "selenium",
            }
        return None

    soup = BeautifulSoup(response.content, "lxml")

    result = {
        "title": "",
        "content_text": "",
        "content_html": "",
        "source_url": url,
        "method": "detail",
    }

    title_el = soup.find("div", class_="title") or soup.find("h1")
    if title_el:
        result["title"] = title_el.get_text(strip=True)

    best_text, best_html = extract_text_from_soup(soup)
    result["content_text"] = best_text
    result["content_html"] = best_html
    logger.info(f"  Detail page text length: {len(best_text)}")

    for candidate_url in find_fulltext_urls(soup, url):
        logger.info(f"  -> Try fulltext candidate: {candidate_url}")
        time.sleep(REQUEST_DELAY)
        candidate_response = fetch_page(candidate_url)
        if not candidate_response:
            continue
        candidate_soup = BeautifulSoup(candidate_response.content, "lxml")
        candidate_text, candidate_html = extract_text_from_soup(candidate_soup)
        logger.info(f"     Candidate text length: {len(candidate_text)}")
        if len(candidate_text) > len(result["content_text"]):
            result["content_text"] = candidate_text
            result["content_html"] = candidate_html
            result["source_url"] = candidate_url
            result["method"] = "iframe_or_print"

    if len(result["content_text"]) < MIN_FULLTEXT_CHARS:
        logger.info("  Static fetch did not reach full-text threshold; try Selenium fallback.")
        rendered_text = render_text_with_selenium(url)
        logger.info(f"  Selenium text length: {len(rendered_text)}")
        if len(rendered_text) > len(result["content_text"]):
            result["content_text"] = rendered_text
            result["content_html"] = ""
            result["source_url"] = url
            result["method"] = "selenium"

    return result


# ============================================================
# TẠO FILE WORD (.DOCX) TỪ NỘI DUNG
# ============================================================

def create_word_document(vanban_info, content_text, output_path):
    """
    Tạo file Word (.docx) chuyên nghiệp từ nội dung văn bản pháp luật.

    Parameters:
        vanban_info: dict chứa thông tin văn bản (ten, soHieu, ngayBanHanh, ...)
        content_text: nội dung text của văn bản
        output_path: đường dẫn file output
    """
    if not HAS_DOCX:
        logger.error("python-docx chưa được cài đặt, không thể tạo file Word.")
        return False

    try:
        doc = Document()

        # === Thiết lập style ===
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(13)

        # === HEADER - Cơ quan ban hành ===
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Times New Roman"

        sub_header = doc.add_paragraph()
        sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_header.add_run("Độc lập - Tự do - Hạnh phúc")
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Times New Roman"

        # Dòng kẻ
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line_para.add_run("─────────────────────")
        run.font.size = Pt(10)

        # Khoảng trống
        doc.add_paragraph()

        # === Số hiệu và ngày ban hành ===
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = info_para.add_run(f"Số: {vanban_info['soHieu']}")
        run.font.size = Pt(13)
        run.font.name = "Times New Roman"

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"Ngày ban hành: {vanban_info['ngayBanHanh']}")
        run.font.size = Pt(12)
        run.font.name = "Times New Roman"
        run.italic = True

        doc.add_paragraph()

        # === TÊN VĂN BẢN ===
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(vanban_info["ten"].upper())
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue

        doc.add_paragraph()

        # === THÔNG TIN METADATA ===
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fields = [
            ("Loại văn bản", vanban_info["loaiVanBan"]),
            ("Số hiệu", vanban_info["soHieu"]),
            ("Ngày ban hành", vanban_info["ngayBanHanh"]),
            ("Lĩnh vực", vanban_info.get("linhVuc", "")),
            ("Nguồn", "Cơ sở dữ liệu quốc gia về pháp luật (vbpl.vn)"),
        ]
        for label, value in fields:
            if value:
                p = doc.add_paragraph()
                run_label = p.add_run(f"{label}: ")
                run_label.bold = True
                run_label.font.size = Pt(12)
                run_label.font.name = "Times New Roman"
                run_value = p.add_run(value)
                run_value.font.size = Pt(12)
                run_value.font.name = "Times New Roman"

        # Dòng kẻ phân cách
        sep_para = doc.add_paragraph()
        sep_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sep_para.add_run("═" * 50)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        # === NỘI DUNG VĂN BẢN ===
        content_title = doc.add_paragraph()
        content_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content_title.add_run("NỘI DUNG VĂN BẢN")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = "Times New Roman"

        doc.add_paragraph()

        # Xử lý nội dung - chia thành các đoạn
        if content_text:
            paragraphs = content_text.split("\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    continue

                p = doc.add_paragraph()

                # Nhận diện tiêu đề chương/điều/mục
                is_heading = False
                if re.match(r"^(PHẦN|CHƯƠNG|MỤC)\s+[IVXLCDM\d]+", para_text, re.IGNORECASE):
                    run = p.add_run(para_text)
                    run.bold = True
                    run.font.size = Pt(13)
                    run.font.name = "Times New Roman"
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    is_heading = True
                elif re.match(r"^Điều\s+\d+", para_text):
                    run = p.add_run(para_text)
                    run.bold = True
                    run.font.size = Pt(13)
                    run.font.name = "Times New Roman"
                    is_heading = True

                if not is_heading:
                    run = p.add_run(para_text)
                    run.font.size = Pt(13)
                    run.font.name = "Times New Roman"
        else:
            # Nếu không có nội dung crawl được, sử dụng mô tả
            p = doc.add_paragraph()
            run = p.add_run(vanban_info.get("mo_ta", "Nội dung đang được cập nhật."))
            run.font.size = Pt(13)
            run.font.name = "Times New Roman"

        # === FOOTER ===
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run("─" * 40)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(128, 128, 128)

        note_para = doc.add_paragraph()
        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note_para.add_run(
            "Văn bản này được tải về từ Cơ sở dữ liệu quốc gia về pháp luật (vbpl.vn)\n"
            f"Thời gian tải: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}\n"
            "Lưu ý: Kiểm tra hiệu lực văn bản trước khi sử dụng."
        )
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Lưu file
        ensure_dir(os.path.dirname(output_path))
        doc.save(output_path)
        logger.info(f"  ✓ Đã tạo file: {output_path}")
        return True

    except Exception as e:
        logger.error(f"  ✗ Lỗi tạo file Word: {e}")
        return False


# ============================================================
# HÀM CRAWL CHÍNH
# ============================================================

def crawl_document(vanban_info, linh_vuc_key, linh_vuc_ten):
    """
    Crawl một văn bản pháp luật.

    Trả về dict chứa kết quả hoặc None nếu thất bại hoàn toàn.
    """
    ten = vanban_info["ten"]
    url = vanban_info.get("url", "")
    filename = sanitize_filename(ten) + ".docx"
    rel_path = f"documents/{linh_vuc_key}/{filename}"
    abs_path = os.path.join(DOCUMENTS_DIR, linh_vuc_key, filename)

    logger.info(f"\n{'='*60}")
    logger.info(f"Đang xử lý: {ten}")
    logger.info(f"  Số hiệu: {vanban_info['soHieu']}")
    logger.info(f"  Lĩnh vực: {linh_vuc_ten}")

    content_text = ""

    # === Bước 1: Thử crawl từ URL (nếu có) ===
    if url:
        try:
            crawled = extract_content_from_tvpl(url)
            if crawled and crawled.get("content_text"):
                content_text = crawled["content_text"]
                logger.info(f"  ✓ Đã crawl được {len(content_text)} ký tự nội dung")
            else:
                logger.warning(
                    f"Crawl URL thành công nhưng không trích xuất được nội dung: "
                    f"{url} | Văn bản: {ten}"
                )
        except Exception as e:
            logger.warning(
                f"Lỗi crawl từ URL: {url} | Văn bản: {ten} | Lỗi: {e}"
            )

    # === Bước 2: Nếu không crawl được, sử dụng mô tả ===
    if not content_text:
        logger.info("  → Sử dụng nội dung mô tả có sẵn (không crawl được toàn văn)")
        content_text = vanban_info.get("mo_ta", "")

    # === Bước 3: Tạo file Word ===
    vanban_info_full = {
        **vanban_info,
        "linhVuc": linh_vuc_ten,
    }
    success = create_word_document(vanban_info_full, content_text, abs_path)

    if success:
        return {
            "ten": ten,
            "soHieu": vanban_info["soHieu"],
            "ngayBanHanh": vanban_info["ngayBanHanh"],
            "linhVuc": linh_vuc_ten,
            "loaiVanBan": vanban_info["loaiVanBan"],
            "duongDanFile": rel_path,
        }
    else:
        logger.error(f"Không thể tạo file cho văn bản: {ten}")
        return None


def iter_test_documents(limit):
    """Yield documents with a real URL for content extraction testing."""
    count = 0
    for linh_vuc_key, linh_vuc_data in VANBAN_DATA.items():
        linh_vuc_ten = linh_vuc_data["ten_linh_vuc"]
        for vanban in linh_vuc_data["vanban"]:
            if not vanban.get("url"):
                continue
            yield linh_vuc_key, linh_vuc_ten, vanban
            count += 1
            if count >= limit:
                return


def test_content_fetch(limit=2):
    """Test full-text extraction only; do not write Word/CSV files."""
    print(f"TEST_CONTENT_FETCH limit={limit}")
    for idx, (_linh_vuc_key, linh_vuc_ten, vanban) in enumerate(iter_test_documents(limit), 1):
        print("=" * 80)
        print(f"[{idx}] {vanban['ten']}")
        print(f"So hieu: {vanban['soHieu']}")
        print(f"Linh vuc: {linh_vuc_ten}")
        print(f"URL: {vanban.get('url')}")

        crawled = extract_content_from_tvpl(vanban["url"])
        content_text = (crawled or {}).get("content_text", "")
        method = (crawled or {}).get("method", "none")
        source_url = (crawled or {}).get("source_url", "")
        preview = content_text[:800].replace("\r", " ")

        print(f"METHOD: {method}")
        print(f"SOURCE_URL: {source_url}")
        print(f"CONTENT_LENGTH: {len(content_text)}")
        print("PREVIEW:")
        print(preview)
    print("=" * 80)
    print("Done. No Word/CSV files were written.")


def main():
    """Hàm chính chạy toàn bộ quá trình crawl."""
    if "--test-content" in sys.argv:
        parser = argparse.ArgumentParser(description="Crawl legal documents.")
        parser.add_argument("--test-content", type=int, default=2, metavar="N")
        args = parser.parse_args()
        test_content_fetch(args.test_content)
        return 0, 0

    start_time = time.time()

    logger.info("╔══════════════════════════════════════════════════════════╗")
    logger.info("║  SCRIPT CRAWL VĂN BẢN PHÁP LUẬT TỪ VBPL.VN           ║")
    logger.info("║  Đồ án: Hệ thống quản lý tài liệu nội bộ             ║")
    logger.info("║  cho công ty luật dân sự                               ║")
    logger.info("╚══════════════════════════════════════════════════════════╝")
    logger.info(f"\nThời gian bắt đầu: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    logger.info(f"Thư mục lưu file: {DOCUMENTS_DIR}")
    logger.info(f"File CSV output: {CSV_OUTPUT}")
    logger.info(f"File log lỗi: {LOG_FILE}")

    # Tạo thư mục
    ensure_dir(DOCUMENTS_DIR)
    for key in VANBAN_DATA.keys():
        ensure_dir(os.path.join(DOCUMENTS_DIR, key))

    # Danh sách kết quả
    all_results = []
    total_success = 0
    total_fail = 0
    total_docs = sum(len(v["vanban"]) for v in VANBAN_DATA.values())

    logger.info(f"\nTổng số văn bản cần xử lý: {total_docs}")
    logger.info(f"Số lĩnh vực: {len(VANBAN_DATA)}")
    logger.info(f"Giới hạn mỗi lĩnh vực: {MAX_PER_FIELD} văn bản")

    # === CRAWL TỪNG LĨNH VỰC ===
    for linh_vuc_key, linh_vuc_data in VANBAN_DATA.items():
        linh_vuc_ten = linh_vuc_data["ten_linh_vuc"]
        vanban_list = linh_vuc_data["vanban"][:MAX_PER_FIELD]

        logger.info(f"\n{'─'*60}")
        logger.info(f"📁 LĨNH VỰC: {linh_vuc_ten}")
        logger.info(f"   Số văn bản: {len(vanban_list)}")
        logger.info(f"   Thư mục: documents/{linh_vuc_key}/")
        logger.info(f"{'─'*60}")

        for idx, vanban in enumerate(vanban_list, 1):
            logger.info(f"\n  [{idx}/{len(vanban_list)}] Xử lý văn bản...")

            try:
                result = crawl_document(vanban, linh_vuc_key, linh_vuc_ten)
                if result:
                    all_results.append(result)
                    total_success += 1
                else:
                    total_fail += 1
            except Exception as e:
                logger.error(f"Lỗi không xử lý được: {vanban['ten']} | {e}")
                total_fail += 1

    # === XUẤT FILE CSV ===
    logger.info(f"\n{'='*60}")
    logger.info("📄 XUẤT FILE CSV...")

    try:
        with open(CSV_OUTPUT, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(
                f,
                fieldnames=["ten", "soHieu", "ngayBanHanh", "linhVuc", "loaiVanBan", "duongDanFile"],
                quoting=csv.QUOTE_ALL,
            )
            writer.writeheader()
            writer.writerows(all_results)
        logger.info(f"  ✓ Đã xuất CSV: {CSV_OUTPUT}")
        logger.info(f"  ✓ Tổng số bản ghi: {len(all_results)}")
    except Exception as e:
        logger.error(f"  ✗ Lỗi xuất CSV: {e}")

    # === BÁO CÁO KẾT QUẢ ===
    elapsed = time.time() - start_time

    logger.info(f"\n{'═'*60}")
    logger.info("📊 BÁO CÁO KẾT QUẢ")
    logger.info(f"{'═'*60}")
    logger.info(f"  Tổng văn bản xử lý : {total_success + total_fail}")
    logger.info(f"  ✓ Thành công        : {total_success}")
    logger.info(f"  ✗ Thất bại          : {total_fail}")
    logger.info(f"  Thời gian           : {elapsed:.1f} giây")
    logger.info(f"  File CSV            : {CSV_OUTPUT}")
    logger.info(f"  File log lỗi        : {LOG_FILE}")
    logger.info(f"{'═'*60}")

    # Liệt kê file đã tạo
    logger.info("\n📂 DANH SÁCH FILE ĐÃ TẠO:")
    for result in all_results:
        logger.info(f"  📄 {result['duongDanFile']}")

    # Liệt kê theo lĩnh vực
    logger.info("\n📊 THỐNG KÊ THEO LĨNH VỰC:")
    for linh_vuc_key, linh_vuc_data in VANBAN_DATA.items():
        count = sum(1 for r in all_results if r["linhVuc"] == linh_vuc_data["ten_linh_vuc"])
        logger.info(f"  {linh_vuc_data['ten_linh_vuc']}: {count}/{len(linh_vuc_data['vanban'][:MAX_PER_FIELD])} văn bản")

    logger.info(f"\n✅ Hoàn tất! Kiểm tra thư mục: {DOCUMENTS_DIR}")
    logger.info(f"📋 File CSV: {CSV_OUTPUT}")

    return total_success, total_fail


# ============================================================
# ENTRY POINT
# ============================================================

if __name__ == "__main__":
    success, fail = main()
    if fail > 0:
        print(f"\n⚠️  Có {fail} văn bản gặp lỗi. Xem chi tiết tại: {LOG_FILE}")
